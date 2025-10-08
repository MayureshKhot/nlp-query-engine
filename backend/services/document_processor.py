import os
import io
import asyncio
from typing import List, Dict, Any
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
import pandas as pd
from core.config import settings

from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy import create_engine
from models.base import Base, Document, Employee
from services.query_processor import QueryProcessor

class DocumentProcessor:
    def __init__(self, db_url: str = settings.DATABASE_URL):
        self.embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL)
        self.vector_store = None
        self.document_metadata = {}
        self.query_processor = QueryProcessor()
        print(f"Initializing DocumentProcessor with database at: {db_url}")
        self.engine = create_engine(db_url)
        Base.metadata.create_all(self.engine)  # Create tables if they don't exist
        print(f"Database tables created. Location: {self.engine.url.database}")
        self._initialize_vector_store()
    
    def _initialize_vector_store(self):
        """Initialize FAISS vector store"""
        dimension = settings.VECTOR_DIMENSION
        self.vector_store = faiss.IndexFlatIP(dimension)  # Inner product for cosine similarity
        print(f"Initialized FAISS vector store with dimension {dimension}")
    
    async def process_files(self, files: List, emp_id: int) -> Dict[str, str]:
        """Process uploaded files and create embeddings"""
        print(f"Processing {len(files)} uploaded files for employee ID {emp_id}...")
        
        processed_data = {}
        all_texts = []
        all_metadata = []
        
        # Create database session
        Session = sessionmaker(bind=self.engine)
        session = Session()
        
        try:
            # Verify employee exists
            employee = session.query(Employee).filter_by(emp_id=emp_id).first()
            if not employee:
                raise ValueError(f"Employee with ID {emp_id} not found")

            for file in files:
                print(f"Processing file: {file.filename}")
                
                try:
                    # Read file content based on type
                    content = await self._extract_text_from_file(file)
                    
                    if content:
                        # Clean up the content
                        content = content.strip()
                        
                        # Determine content type
                        content_type = 'resume' if 'resume' in file.filename.lower() else 'general'
                        
                        # Store in database
                        document = Document(
                            filename=file.filename,
                            content=content,
                            content_type=content_type,
                            emp_id=emp_id
                        )
                        
                        # Check if document already exists
                        existing_doc = session.query(Document).filter_by(filename=file.filename, emp_id=emp_id).first()
                        if existing_doc:
                            existing_doc.content = content
                            existing_doc.content_type = content_type
                        else:
                            session.add(document)
                        
                        session.commit()
                        processed_data[file.filename] = content
                        
                    print("\n" + "="*50)
                    print(f"Document Processing Details for: {file.filename}")
                    print("="*50)
                    print(f"Total content length: {len(content)} characters")
                    print("\nOriginal Content Preview:")
                    print("-"*30)
                    print(content[:500] + "..." if len(content) > 500 else content)
                    print("-"*30)
                    
                    # Split content into chunks for better search
                    chunks = self._split_text_into_chunks(content, file.filename)
                    
                    print(f"\nCreated {len(chunks)} chunks for processing")
                    print("\nChunk Details:")
                    print("-"*30)
                    
                    for i, chunk in enumerate(chunks):
                        if chunk.strip():  # Only process non-empty chunks
                            all_texts.append(chunk)
                            metadata = {
                                "filename": file.filename,
                                "chunk_id": i,
                                "chunk_length": len(chunk),
                                "content": chunk,
                                "document_id": document.id  # Store the database ID
                            }
                            all_metadata.append(metadata)
                            
                            print(f"\nChunk {i}:")
                            print(f"Length: {len(chunk)} characters")
                            print(f"Content:\n{chunk}")
                            print("-"*30)                        
                            print(f"-> Successfully processed {file.filename} into {len(chunks)} chunks")
                        
                except Exception as e:
                    print(f"!! Error processing file {file.filename}: {e}")
                    processed_data[file.filename] = f"Error: Could not process file - {str(e)}"
                    
            # Create embeddings for all chunks
            if all_texts:
                await self._create_embeddings(all_texts, all_metadata)
            
            session.commit()
            return processed_data
            
        except Exception as e:
            session.rollback()
            raise e
        
        finally:
            session.close()
    
    async def _extract_text_from_file(self, file) -> str:
        """Extract text from different file formats"""
        content_bytes = file.file.read()
        file.file.seek(0)  # Reset file pointer
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                return self._extract_from_pdf(content_bytes)
            elif filename.endswith('.docx'):
                return self._extract_from_docx(content_bytes)
            elif filename.endswith('.txt'):
                return content_bytes.decode('utf-8')
            elif filename.endswith('.csv'):
                return self._extract_from_csv(content_bytes)
            else:
                # Try to decode as text
                return content_bytes.decode('utf-8')
        except Exception as e:
            print(f"Error extracting text from {file.filename}: {e}")
            return ""
    
    def _extract_from_pdf(self, content_bytes: bytes) -> str:
        """Extract text from PDF"""
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_from_docx(self, content_bytes: bytes) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(io.BytesIO(content_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_csv(self, content_bytes: bytes) -> str:
        """Extract text from CSV"""
        df = pd.read_csv(io.BytesIO(content_bytes))
        return df.to_string()
    
    def _split_text_into_chunks(self, text: str, filename: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks for better search"""
        print(f"Splitting text from {filename} into chunks (length: {len(text)})")
        
        # Always store the complete document as the first chunk
        chunks = [text] if text else []
        print("Added complete document as first chunk")
        
        # For resumes or short documents, just return the complete document
        if len(text) < 2000 or "resume" in filename.lower():
            print(f"Document is a resume or short text, keeping as single chunk ({len(text)} chars)")
            return chunks
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for paragraph breaks first
                last_para = chunk.rfind('\n\n')
                if last_para > start + chunk_size // 2:
                    chunk = chunk[:last_para].strip()
                    end = start + last_para
                else:
                    # Try sentence breaks
                    sentence_breaks = ['. ', '? ', '! ', '.\n', '?\n', '!\n']
                    break_points = [chunk.rfind(sep) for sep in sentence_breaks]
                    break_points = [bp for bp in break_points if bp > start + chunk_size // 2]
                    
                    if break_points:
                        break_point = max(break_points)
                        chunk = chunk[:break_point + 1].strip()
                        end = start + break_point + 1
            
            # Clean up the chunk
            chunk = chunk.strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
                print(f"Created chunk {len(chunks)}: {len(chunk)} chars")
            
            start = end - overlap
        
        print(f"Split into {len(chunks)} chunks")
        return chunks
    
    async def _create_embeddings(self, texts: List[str], metadata: List[Dict]):
        """Create embeddings for text chunks and add to vector store"""
        print("\n" + "="*50)
        print("Creating Vector Embeddings")
        print("="*50)
        print(f"Processing {len(texts)} text chunks for vector embeddings...")
        
        try:
            # Generate embeddings
            print("\nGenerating embeddings with model:", settings.EMBEDDING_MODEL)
            embeddings = self.embedding_model.encode(texts, show_progress_bar=True)
            
            # Normalize embeddings for cosine similarity
            print("\nNormalizing embeddings for improved similarity matching...")
            norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
            norms[norms == 0] = 1  # Prevent division by zero
            embeddings = embeddings / norms
            
            print(f"\nEmbedding Details:")
            print(f"- Shape: {embeddings.shape}")
            print(f"- Dimension: {embeddings.shape[1]}")
            print(f"- Number of chunks: {embeddings.shape[0]}")
            
            # Clear existing index and create new one
            self._initialize_vector_store()
            
            # Add embeddings to FAISS index
            self.vector_store.add(embeddings.astype('float32'))
            
            # Store metadata with correct indices
            self.document_metadata.clear()  # Clear existing metadata
            for i, meta in enumerate(metadata):
                self.document_metadata[i] = meta
                
            print(f"Successfully added {len(texts)} embeddings to vector store")
            print(f"Vector store total: {self.vector_store.ntotal}")
            print(f"Metadata entries: {len(self.document_metadata)}")
            
        except Exception as e:
            print(f"Error creating embeddings: {str(e)}")
            raise
        
        print(f"-> Added {len(texts)} embeddings to vector store. Total vectors: {self.vector_store.ntotal}")
    
    def search_similar_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents using semantic similarity"""
        if self.vector_store is None or self.vector_store.ntotal == 0:
            print("No documents in vector store")
            return []
        
        print(f"Searching for: '{query}' in {self.vector_store.ntotal} vectors")
        print(f"Current document metadata: {list(self.document_metadata.keys())}")
        
        # Create database session
        Session = sessionmaker(bind=self.engine)
        session = Session()
        
        # Clean and prepare query
        query = query.lower().strip()
        # Extract key terms (names, skills, etc.)
        key_terms = [term.strip() for term in query.split() if len(term.strip()) > 2]
        search_query = " ".join(key_terms)
        
        # Create embedding for query
        try:
            query_embedding = self.embedding_model.encode([query])
            query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
            
            # Search in vector store
            scores, indices = self.vector_store.search(query_embedding.astype('float32'), min(top_k, self.vector_store.ntotal))
            
            try:
                results = []
                print("\n" + "="*50)
                print(f"Search Results")
                print("="*50)
                print(f"Found {len(indices[0])} potential matches")
                
                for score, idx in zip(scores[0], indices[0]):
                    if idx in self.document_metadata:
                        metadata = self.document_metadata[idx]
                        document = session.query(Document).filter_by(filename=metadata["filename"]).first()
                        
                        if document:
                            chunk_content = metadata["content"]
                            print("\n" + "-"*30)
                            print(f"Match {len(results) + 1}:")
                            print(f"Document: {document.filename}")
                            print(f"Content Type: {document.content_type}")
                            print(f"Similarity Score: {score:.4f}")
                            print(f"Chunk ID: {metadata['chunk_id']}")
                            print("\nMatched Content:")
                            print(chunk_content)
                            print("-"*30)
                            
                            # Format the preview to show context around matching terms
                            preview = chunk_content
                            if len(preview) > 300:
                                preview = preview[:300] + "..."
                            
                            result = {
                                "source": document.filename,
                                "document_id": document.id,
                                "chunk_id": metadata["chunk_id"],
                                "similarity_score": float(score),
                                "content_preview": preview,
                                "full_document": document.content,
                                "content_type": document.content_type,
                                "last_updated": document.updated_at.isoformat()
                            }
                            
                            print(f"Match in {document.filename} (score: {score:.3f})")
                            results.append(result)
                
            finally:
                session.close()
            
            return results
            
        except Exception as e:
            print(f"Error during document search: {str(e)}")
            return []
    
    def get_document_count(self) -> int:
        """Get total number of documents in the store"""
        return self.vector_store.ntotal if self.vector_store else 0
    
    def clear_documents(self):
        """Clear all documents and embeddings"""
        self._initialize_vector_store()
        self.document_metadata = {}
        print("Document store cleared")
