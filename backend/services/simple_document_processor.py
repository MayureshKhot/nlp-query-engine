import os
import io
import time
from typing import List, Dict, Any
import PyPDF2
import docx
import pandas as pd
from core.config import settings
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
import io
from typing import List, Dict, Any
import PyPDF2
import docx
import pandas as pd
from core.config import settings
from sqlalchemy import create_engine, or_
from sqlalchemy.orm import sessionmaker, Session
from models.base import Base, Document, Employee, Department

class SimpleDocumentProcessor:
    def __init__(self, db_url: str = settings.DATABASE_URL):
        self.engine = create_engine(db_url)
        Base.metadata.create_all(self.engine)
        self.Session = sessionmaker(bind=self.engine)
        print("Simple Document Processor initialized with database")
    
    async def process_files(self, files: List, emp_id: int) -> Dict[str, str]:
        """Process uploaded files without embeddings"""
        print(f"Processing {len(files)} uploaded files for employee {emp_id}...")
        
        processed_data = {}
        session = self.Session()
        
        try:
            for file in files:
                print(f"Processing file: {file.filename}")
                
                try:
                    # Read file content based on type
                    content = await self._extract_text_from_file(file)
                    
                    if content:
                        # Determine content type
                        content_type = 'resume' if 'resume' in file.filename.lower() else 'general'
                        
                        # Create or update document in database
                        existing_doc = session.query(Document).filter_by(
                            filename=file.filename,
                            emp_id=emp_id
                        ).first()
                        
                        if existing_doc:
                            existing_doc.content = content
                            existing_doc.content_type = content_type
                        else:
                            new_doc = Document(
                                filename=file.filename,
                                content=content,
                                content_type=content_type,
                                emp_id=emp_id
                            )
                            session.add(new_doc)
                        
                        processed_data[file.filename] = content
                        print(f"-> Successfully processed and stored {file.filename}")
                    else:
                        print(f"-> No content extracted from {file.filename}")
                        
                except Exception as e:
                    print(f"!! Error processing file {file.filename}: {e}")
                    processed_data[file.filename] = f"Error: Could not process file - {str(e)}"
            
            session.commit()
            return processed_data
            
        except Exception as e:
            session.rollback()
            raise e
        
        finally:
            session.close()
    
    async def _extract_text_from_file(self, file) -> str:
        """Extract text from different file formats"""
        content_bytes = file.file.read()
        file.file.seek(0)  # Reset file pointer
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                return self._extract_from_pdf(content_bytes)
            elif filename.endswith('.docx'):
                return self._extract_from_docx(content_bytes)
            elif filename.endswith('.txt'):
                return content_bytes.decode('utf-8')
            elif filename.endswith('.csv'):
                return self._extract_from_csv(content_bytes)
            else:
                # Try to decode as text
                return content_bytes.decode('utf-8')
        except Exception as e:
            print(f"Error extracting text from {file.filename}: {e}")
            return ""
    
    def _extract_from_pdf(self, content_bytes: bytes) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_from_docx(self, content_bytes: bytes) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(io.BytesIO(content_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_csv(self, content_bytes: bytes) -> str:
        """Extract text from CSV"""
        df = pd.read_csv(io.BytesIO(content_bytes))
        return df.to_string()
    
    def search_similar_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for documents based on content and metadata"""
        print(f"Searching for: '{query}'")
        
        session = self.Session()
        try:
            # Clean and prepare query
            query = query.lower().strip()
            start_time = time.time()  # For calculating similarity score

            # Immediately return empty list for empty query
            if not query.strip():
                return []
            
            # Handle special queries
            if any(phrase in query for phrase in [
                "show all documents", 
                "available documents",
                "list documents",
                "show documents"
            ]):
                documents = session.query(Document).all()
                return [self._format_document_result(doc, query, session, 1.0) for doc in documents[:top_k]]

            # Break query into search terms
            search_terms = query.split()
            base_query = session.query(Document)

            # Handle resume-specific searches
            if "resume" in query:
                # Extract potential name (words before "resume")
                resume_idx = search_terms.index("resume")
                potential_name = " ".join(search_terms[:resume_idx])
                if potential_name:
                    # Search by employee name
                    documents = base_query.join(Employee).filter(
                        Employee.full_name.ilike(f"%{potential_name}%")
                    ).all()
                    if documents:
                        return [self._format_document_result(doc, query, session, 0.9) for doc in documents[:top_k]]

            # General content search
            conditions = []
            for term in search_terms:
                if len(term) > 2:  # Only search for terms longer than 2 characters
                    conditions.append(Document.content.ilike(f"%{term}%"))
                    conditions.append(Document.filename.ilike(f"%{term}%"))
                    # Also search in employee names
                    conditions.append(Employee.full_name.ilike(f"%{term}%"))

            if conditions:
                documents = base_query.join(Employee).filter(or_(*conditions)).all()
                return [self._format_document_result(doc, query, session, 0.8) for doc in documents[:top_k]]

            return []

        finally:
            session.close()

    def _format_document_result(self, doc: Document, query: str, session: Session, base_score: float) -> Dict[str, Any]:
        """Format a single document result for the frontend"""
        try:
            # Get employee details
            employee = session.query(Employee).filter_by(emp_id=doc.emp_id).first()
            
            # Create content preview focusing on query terms
            preview = doc.content
            if len(preview) > 300:
                # Try to center preview around first query term match
                query_terms = [term for term in query.lower().split() if len(term) > 2]
                for term in query_terms:
                    if term in doc.content.lower():
                        idx = doc.content.lower().find(term)
                        start = max(0, idx - 150)
                        end = min(len(doc.content), idx + 150)
                        preview = "..." + doc.content[start:end] + "..."
                        break
                else:
                    preview = doc.content[:300] + "..."

            # Calculate content relevance score
            score = base_score
            if query_terms := set(query.lower().split()):
                content_terms = set(doc.content.lower().split())
                matching_terms = query_terms.intersection(content_terms)
                if matching_terms:
                    term_score = len(matching_terms) / len(query_terms) * 0.2
                    score = min(score + term_score, 1.0)

            result = {
                "source": doc.filename,  # Used as title in frontend
                "similarity_score": score,
                "content_preview": preview,
                "content_type": doc.content_type,
                "full_document": doc.content,
                "last_updated": doc.updated_at.isoformat() if doc.updated_at else None
            }

            # Add employee information if available
            if employee:
                result["employee"] = {
                    "id": employee.emp_id,
                    "name": employee.full_name,
                    "department": employee.department.dept_name if employee.department else None
                }

            return result

        except Exception as e:
            print(f"Error formatting document result: {str(e)}")
            # Return a minimal result in case of error
            return {
                "source": doc.filename,
                "similarity_score": 0.0,
                "content_preview": "Error processing document",
                "content_type": "unknown",
                "full_document": "",
                "last_updated": None
            }
            
    def _format_search_results(self, documents: List[Document], query: str, session: Session) -> List[Dict[str, Any]]:
        """Format document search results for the frontend"""
        results = []
        query_terms = query.lower().split()
        
        for doc in documents:
            # Get employee details
            employee = session.query(Employee).filter_by(emp_id=doc.emp_id).first()
            
            # Create content preview focusing on query terms
            preview = doc.content
            if len(preview) > 300:
                # Try to center preview around first query term match
                for term in query_terms:
                    if term in doc.content.lower():
                        idx = doc.content.lower().find(term)
                        start = max(0, idx - 150)
                        end = min(len(doc.content), idx + 150)
                        preview = "..." + doc.content[start:end] + "..."
                        break
                else:
                    preview = doc.content[:300] + "..."
            
            result = {
                "source": doc.filename,
                "document_id": doc.id,
                "content_preview": preview,
                "full_document": doc.content,
                "content_type": doc.content_type,
                "last_updated": doc.updated_at.isoformat(),
                "employee": {
                    "id": employee.emp_id,
                    "name": employee.full_name,
                    "department": employee.department.dept_name if employee.department else None
                } if employee else None
            }
            results.append(result)
        
        return results
    
    def get_document_count(self) -> int:
        """Get total number of documents in the store"""
        session = self.Session()
        try:
            return session.query(Document).count()
        finally:
            session.close()
    
    def clear_documents(self):
        """Clear all documents"""
        session = self.Session()
        try:
            session.query(Document).delete()
            session.commit()
            print("Document store cleared")
        except Exception as e:
            session.rollback()
            raise e
        finally:
            session.close()
